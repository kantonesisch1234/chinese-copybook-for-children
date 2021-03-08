import docx 
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn
from docx.shared import RGBColor
from google_trans_new import google_translator
from docx2pdf import convert
import requests
from bs4 import BeautifulSoup
import json
import os
import sys
import shutil

black = RGBColor(0, 0, 0)
gray = RGBColor(220, 220, 220)

# A convenient function to set English and Chinese fonts
def set_run_font(run,size,english_font='Times New Roman',chinese_font=u'標楷體',font_color=black):
    font = run.font
    font.name = english_font
    font.size = Pt(size)
    font.color.rgb = font_color
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)

class copybook_page:
    def __init__(self, word, english_word, mode, filename="copybook.docx"):
        assert len(word) <= 10
        self.word = word
        self.english_word = english_word
        self.mode = mode       # input from command line, "-t" is translation mode and "-p" is picture inserting mode  
        self.filename = filename
        
    def __download_image_from_word(self):
        url = "https://images.search.yahoo.com/search/images?p=" + self.word

        response = requests.get(url)
        soup = BeautifulSoup(response.text, "html.parser")
        img_link = json.loads(soup.find("div", class_="sres-cntr").find("li")['data'])['iurl']
        
        if not os.path.exists('pics'):
            os.makedirs('pics')

        response = requests.get(img_link)
        file = open("pics/"+self.word+".jpg" , "wb")
        file.write(response.content)
        file.close()
    
    def insert_to_document(self, document):
        wordlen = len(self.word)
        
        cell_dim_list = [(8,7), (10,8), (8,6), (10,8), (14,10), (8,6), 
                         (8,7), (10,8), (10,9), (14,10)]
        cell_size_list = [(2.5,2.5), (2,2), (2.5,2.5), (2,2), (1.5,1.5), (2.5,2.5), (2.5,2.5), (2,2), (2,2), (1.5,1.5)]
        title_font_size_list = [36, 36, 36, 36, 32, 32, 24, 24, 24, 20]
        translation_font_size_list = [16,16,16,12,12,12,10,10,10,10]
        font_size_list = [52, 42, 52, 42, 32, 52, 52, 42, 36, 36]
        merge_cell_list = [[(0,0),(1,2),(0,3),(1,6)],
                          [(0,0),(1,3),(0,4),(1,7)],
                          [(0,0),(1,2),(0,3),(1,5)],
                          [(0,0),(1,3),(0,4),(1,7)],
                          [(0,0),(1,4),(0,5),(1,9)],
                          [(0,0),(1,2),(0,3),(1,5)],
                          [(0,0),(1,2),(0,3),(1,6)],
                          [(0,0),(1,3),(0,4),(1,7)],
                          [(0,0),(1,4),(0,5),(1,8)],
                          [(0,0),(1,4),(0,5),(1,9)]]

        cell_dim = cell_dim_list[wordlen-1]
        cell_size = cell_size_list[wordlen-1]
        title_font_size = title_font_size_list[wordlen-1]
        font_size = font_size_list[wordlen-1]
        
        pic_width = 2.5

        cell_height, cell_width = cell_size
        row_no, column_no = cell_dim
                
        table = document.add_table(rows=row_no, cols=column_no, style='Table Grid')

        for row in table.rows:
            row.height = Cm(cell_height)
            row.width = Cm(cell_width)
            
        # Merge cells
        a = table.cell(merge_cell_list[wordlen-1][0][0], merge_cell_list[wordlen-1][0][1])
        b = table.cell(merge_cell_list[wordlen-1][1][0], merge_cell_list[wordlen-1][1][1])
        a.merge(b)
        
        c = table.cell(merge_cell_list[wordlen-1][2][0], merge_cell_list[wordlen-1][2][1])
        d = table.cell(merge_cell_list[wordlen-1][3][0], merge_cell_list[wordlen-1][3][1])
        c.merge(d)
        
        # Write in the first cell
        title_cell = table.rows[0].cells[1]
        title_cell.text = self.word
        paragraphs = title_cell.paragraphs
        paragraphs[0].alignment = 1
        paragraphs[0].add_run("\n"+self.english_word)
        for paragraph in paragraphs:
            for idx, run in enumerate(paragraph.runs):
                set_run_font(run,title_font_size)
                if idx==1:
                    run.italic = True

        if self.mode == '-t':
            translator = google_translator()
            title_cell_2 = table.rows[0].cells[column_no-1]
            translation = '德文：'+translator.translate(self.english_word,lang_src='en',lang_tgt='de').strip(' ')+'\n'
            translation += '法文：'+translator.translate(self.english_word,lang_src='en',lang_tgt='fr').strip(' ')+'\n'
            translation += '日文：'+translator.translate(self.english_word,lang_src='en',lang_tgt='ja').strip(' ')+'\n'
            translation += '俄文：'+translator.translate(self.english_word,lang_src='en',lang_tgt='ru').strip(' ')
            title_cell_2.text = translation
            paragraphs = title_cell_2.paragraphs
            paragraphs[0].alignment = 1
            run = paragraphs[0].runs[0]
            set_run_font(run,translation_font_size_list[wordlen-1])
        
        if self.mode == '-p':
            title_cell_2 = table.rows[0].cells[column_no-1]
            paragraphs = title_cell_2.paragraphs
            run = paragraphs[0].add_run()
            paragraphs[0].alignment = 1
            run = paragraphs[0].runs[0]
            
            self.__download_image_from_word()
            pic_dir = "pics/"+self.word+".jpg"
            if not os.path.exists(pic_dir):
                return 0    # skip inserting the picture if picture download not successful
            inline_shape = run.add_picture(pic_dir,width=Inches(pic_width))
                
        row = table.rows[2]

        for idx,char in enumerate(self.word):
            row.cells[idx].text = char
        
        for row in table.rows[2:]:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run,font_size,font_color=gray)
                        
        document.save(self.filename)
        document.add_page_break()
        
def create_copybook_from_txt():
    with open("wordlist.txt", 'r', encoding='utf-8') as f:
        lines = f.readlines()

    chinese_words = []
    english_words = []
    for line in lines:
        word_pair = line.strip('\n').split(',')
        chinese_words.append(word_pair[0].strip(' '))
        english_words.append(word_pair[1].strip(' '))

    doc = docx.Document()
    copybook_obj_list = []
    for words in zip(chinese_words, english_words):
        copybook_page(*words,sys.argv[1]).insert_to_document(doc)

    convert("copybook.docx")
    
if __name__ == '__main__':
    create_copybook_from_txt()
    shutil.rmtree("pics")